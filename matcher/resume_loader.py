"""Resume loading.

Reads PDF / Markdown / DOCX / TXT into a single plain-text string.
We deliberately do NOT structure-extract (per Stage 2 design decision Q4+Q6:
zero LLM dependency; structure goes into preferences.yaml instead).

The plain text is the input to the embedding model — semantic similarity
is the only resume signal that flows into scoring. All hard filters and
keyword matching come from `config/preferences.yaml`.
"""
from __future__ import annotations

import re
from pathlib import Path

from pipeline.clean_jd import normalize_whitespace, strip_html


SUPPORTED_SUFFIXES = {".md", ".markdown", ".txt", ".pdf", ".docx"}


class ResumeLoadError(Exception):
    pass


def _load_pdf(path: Path) -> str:
    try:
        import pdfplumber  # type: ignore
    except ImportError as e:
        raise ResumeLoadError(
            "pdfplumber is required for PDF resumes. "
            "Run: pip install pdfplumber>=0.11.0"
        ) from e
    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            if t:
                parts.append(t)
    return "\n\n".join(parts)


def _load_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise ResumeLoadError(
            "python-docx is required for DOCX resumes. "
            "Run: pip install python-docx>=1.1.0"
        ) from e
    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def _load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def load_resume(path: str | Path, min_chars: int = 100) -> str:
    """Read a resume file into a single cleaned plain-text string.

    Args:
      path:      file path. Suffix decides loader.
      min_chars: reject resumes shorter than this — usually means parsing failed
                 (e.g. PDF with image-only content / OCR not done).

    Returns:
      Plain text. Whitespace normalized. HTML tags stripped if any.
    """
    p = Path(path)
    if not p.exists():
        raise ResumeLoadError(f"Resume not found: {p}")
    suffix = p.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ResumeLoadError(
            f"Unsupported resume format '{suffix}'. "
            f"Supported: {sorted(SUPPORTED_SUFFIXES)}"
        )

    if suffix == ".pdf":
        raw = _load_pdf(p)
    elif suffix == ".docx":
        raw = _load_docx(p)
    else:
        raw = _load_text(p)

    text = strip_html(raw)
    text = normalize_whitespace(text)
    text = re.sub(r"[ \t]{2,}", " ", text)

    if len(text) < min_chars:
        raise ResumeLoadError(
            f"Resume parsed to only {len(text)} chars (min={min_chars}). "
            f"File: {p}. If this is a scanned PDF, OCR is required first."
        )

    return text
